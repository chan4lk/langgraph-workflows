from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
import time
import re
from markdown_it import MarkdownIt
from markdown_it.token import Token

def detect_and_parse_pipe_tables(markdown_text):
    """Detect and parse pipe-style markdown tables.
    
    Returns a list of dictionaries with table information:
    {
        'start_line': int,
        'end_line': int,
        'headers': list of strings,
        'rows': list of lists of strings
    }
    """
    tables = []
    lines = markdown_text.split('\n')
    
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        
        # Look for a line that starts with | and has multiple | characters
        if line.startswith('|') and line.count('|') > 1:
            # This might be a table header
            potential_header = line
            headers = [cell.strip() for cell in potential_header.split('|')]
            headers = [h for h in headers if h]  # Remove empty cells
            
            # Check if the next line is a separator line (contains only |, -, and :)
            if i + 1 < len(lines) and all(c in '|:-' for c in lines[i + 1].strip()):
                # This is a table! Extract it
                table_data = {
                    'start_line': i,
                    'headers': headers,
                    'rows': []
                }
                
                # Skip the separator line
                i += 2
                
                # Extract all rows until we hit a non-table line
                while i < len(lines) and lines[i].strip().startswith('|'):
                    row_line = lines[i].strip()
                    row_cells = [cell.strip() for cell in row_line.split('|')]
                    row_cells = [cell for cell in row_cells if cell != '']  # Remove empty cells
                    table_data['rows'].append(row_cells)
                    i += 1
                
                table_data['end_line'] = i - 1
                tables.append(table_data)
                continue  # Skip the increment at the end of the loop
        i += 1
    
    return tables

def add_table_to_doc(doc, table_data):
    """Add a table to the Word document based on parsed table data."""
    if not table_data['headers'] and not table_data['rows']:
        return
    
    # Determine the number of columns
    num_cols = max(
        len(table_data['headers']),
        max([len(row) for row in table_data['rows']]) if table_data['rows'] else 0
    )
    
    # Create the table
    table = doc.add_table(rows=1, cols=num_cols)
    table.style = 'Table Grid'
    table.autofit = False
    
    # Add headers
    header_cells = table.rows[0].cells
    for i, header in enumerate(table_data['headers']):
        if i < len(header_cells):
            header_cells[i].text = ""
            p = header_cells[i].paragraphs[0]
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = p.add_run(header)
            run.bold = True
            
            # Add light gray shading to header cells
            tc = header_cells[i]._tc
            tcPr = tc.get_or_add_tcPr()
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), 'E6E6E6')  # Light gray
            tcPr.append(shading_elm)
    
    # Add data rows
    for row_data in table_data['rows']:
        row_cells = table.add_row().cells
        for i, cell_text in enumerate(row_data):
            if i < len(row_cells):
                row_cells[i].text = cell_text

def markdown_to_docx(doc, markdown_text):
    """Convert markdown text to properly formatted Word document."""
    if not markdown_text:
        return
    
    # First, detect and extract pipe tables
    tables = detect_and_parse_pipe_tables(markdown_text)
    
    # Replace table sections with placeholders to avoid processing them with markdown-it
    placeholder_text = "TABLE_PLACEHOLDER_{}"
    modified_text = markdown_text
    for i, table in enumerate(tables):
        # Extract the table text
        table_lines = markdown_text.split('\n')[table['start_line']:table['end_line']+1]
        table_text = '\n'.join(table_lines)
        
        # Replace with placeholder
        modified_text = modified_text.replace(table_text, placeholder_text.format(i))
    
    # Parse the modified markdown (without tables)
    md = MarkdownIt()
    tokens = md.parse(modified_text)
    
    # Track list types and levels
    list_types = []
    list_level = 0
    
    # Track table state for markdown-it tables (different from pipe tables)
    in_table = False
    table = None
    current_row = None
    header_processed = False
    col_count = 0
    
    # Process tokens
    i = 0
    while i < len(tokens):
        token = tokens[i]
        
        # Check for table placeholders
        if token.type == 'paragraph_open':
            i += 1  # Move to content token
            if i < len(tokens) and tokens[i].content.startswith("TABLE_PLACEHOLDER_"):
                # Extract table index from placeholder
                try:
                    table_idx = int(tokens[i].content.replace("TABLE_PLACEHOLDER_", ""))
                    if table_idx < len(tables):
                        # Add the actual table
                        add_table_to_doc(doc, tables[table_idx])
                except ValueError:
                    # Not a valid placeholder, treat as normal paragraph
                    pass
                
                # Skip to end of paragraph
                while i < len(tokens) and tokens[i].type != 'paragraph_close':
                    i += 1
                i += 1
                continue
        
        # Handle headings
        if token.type == 'heading_open':
            level = int(token.tag[1])  # h1, h2, etc.
            i += 1  # Move to content token
            content = tokens[i].content if i < len(tokens) else ""
            doc.add_heading(content, level=level)
            i += 2  # Skip closing tag
            continue
            
        # Handle paragraphs
        elif token.type == 'paragraph_open':
            i += 1  # Move to content token
            if i >= len(tokens):
                i += 1
                continue
                
            content = tokens[i].content
            p = doc.add_paragraph()
            
            # Process inline formatting
            parts = re.split(r'(\*\*|\*|__|_|`)', content)
            is_bold = False
            is_italic = False
            is_code = False
            
            for part in parts:
                if part == '**' or part == '__':
                    is_bold = not is_bold
                elif part == '*' or part == '_':
                    is_italic = not is_italic
                elif part == '`':
                    is_code = not is_code
                else:
                    run = p.add_run(part)
                    run.bold = is_bold
                    run.italic = is_italic
                    if is_code:
                        run.font.name = 'Courier New'
                        run.font.size = Pt(10)
            
            i += 2  # Skip closing tag
            continue
            
        # Handle lists
        elif token.type == 'bullet_list_open':
            list_types.append('bullet')
            list_level += 1
            i += 1
            continue
        elif token.type == 'ordered_list_open':
            list_types.append('ordered')
            list_level += 1
            i += 1
            continue
        elif token.type == 'bullet_list_close' or token.type == 'ordered_list_close':
            if list_types:
                list_types.pop()
            list_level = max(0, list_level - 1)
            i += 1
            continue
        elif token.type == 'list_item_open':
            i += 1
            # Find the content of this list item
            content = ""
            while i < len(tokens) and tokens[i].type != 'list_item_close':
                if tokens[i].type == 'paragraph_open':
                    i += 1
                    if i < len(tokens):
                        content = tokens[i].content
                    i += 1
                else:
                    i += 1
            
            # Add the list item with proper indentation
            p = doc.add_paragraph()
            
            # Create bullet points manually with proper indentation
            indent = list_level * 0.25
            p.paragraph_format.left_indent = Inches(indent)
            p.paragraph_format.first_line_indent = Inches(-0.25)
            
            # Add bullet character or number based on list type
            current_list_type = list_types[-1] if list_types else 'bullet'
            if current_list_type == 'bullet':
                bullet_char = '•' if list_level % 2 == 1 else '○'
                p.add_run(f"{bullet_char} ").bold = True
            else:  # ordered list
                p.add_run(f"1. ").bold = True
                
            # Add the content with any formatting
            parts = re.split(r'(\*\*|\*|__|_|`)', content)
            is_bold = False
            is_italic = False
            is_code = False
            
            for part in parts:
                if part == '**' or part == '__':
                    is_bold = not is_bold
                elif part == '*' or part == '_':
                    is_italic = not is_italic
                elif part == '`':
                    is_code = not is_code
                else:
                    run = p.add_run(part)
                    run.bold = is_bold
                    run.italic = is_italic
                    if is_code:
                        run.font.name = 'Courier New'
                        run.font.size = Pt(10)
            
            i += 1
            continue
            
        # Handle code blocks
        elif token.type == 'fence' and token.tag == 'code':
            # Add a code block with monospaced font and gray background
            p = doc.add_paragraph()
            code_text = token.content
            
            # Add light gray shading to the paragraph
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), 'F5F5F5')  # Light gray
            p._p.get_or_add_pPr().append(shading_elm)
            
            # Format code with monospaced font
            code_run = p.add_run(code_text)
            code_run.font.name = 'Courier New'
            code_run.font.size = Pt(10)
            
            i += 1
            continue
            
        # Handle horizontal rules
        elif token.type == 'hr':
            p = doc.add_paragraph()
            p.add_run('_' * 50)
            i += 1
            continue
        
        # Handle tables
        elif token.type == 'table_open':
            in_table = True
            # First, determine the number of columns by scanning ahead
            col_count = 0
            temp_idx = i + 1
            while temp_idx < len(tokens) and tokens[temp_idx].type != 'table_close':
                if tokens[temp_idx].type == 'tr_open':
                    # Count cells in this row
                    cell_count = 0
                    inner_idx = temp_idx + 1
                    while inner_idx < len(tokens) and tokens[inner_idx].type != 'tr_close':
                        if tokens[inner_idx].type in ['th_open', 'td_open']:
                            cell_count += 1
                        inner_idx += 1
                    col_count = max(col_count, cell_count)
                    break
                temp_idx += 1
            
            # Create the table with the determined number of columns
            if col_count > 0:
                table = doc.add_table(rows=0, cols=col_count)
                table.style = 'Table Grid'
                # Make the table take up the full width
                table.autofit = False
                table.allow_autofit = False
            else:
                # Fallback if we couldn't determine columns
                table = doc.add_table(rows=0, cols=1)
                table.style = 'Table Grid'
            
            i += 1
            continue
        elif token.type == 'thead_open':
            header_processed = False
            i += 1
            continue
        elif token.type == 'tr_open':
            # Add a row to the table
            if table is not None:
                current_row = table.add_row().cells
            i += 1
            continue
        elif token.type in ['th_open', 'td_open']:
            cell_idx = 0
            # Find which cell we're in
            temp_idx = i
            while temp_idx > 0 and tokens[temp_idx].type != 'tr_open':
                if tokens[temp_idx].type in ['th_close', 'td_close']:
                    cell_idx += 1
                temp_idx -= 1
            
            # Move to content token
            i += 1
            if i < len(tokens) and current_row is not None and cell_idx < len(current_row):
                cell_text = tokens[i].content
                
                # Format header cells with bold text and center alignment
                if token.type == 'th_open':
                    current_row[cell_idx].text = ""
                    p = current_row[cell_idx].paragraphs[0]
                    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    run = p.add_run(cell_text)
                    run.bold = True
                    
                    # Add light gray shading to header cells
                    tc = current_row[cell_idx]._tc
                    tcPr = tc.get_or_add_tcPr()
                    shading_elm = OxmlElement('w:shd')
                    shading_elm.set(qn('w:fill'), 'E6E6E6')  # Light gray
                    tcPr.append(shading_elm)
                else:
                    current_row[cell_idx].text = cell_text
            
            # Skip to closing tag
            while i < len(tokens) and tokens[i].type not in ['th_close', 'td_close']:
                i += 1
            i += 1
            continue
        elif token.type == 'tr_close':
            i += 1
            continue
        elif token.type in ['thead_close', 'tbody_open', 'tbody_close']:
            i += 1
            continue
        elif token.type == 'table_close':
            in_table = False
            table = None
            current_row = None
            header_processed = False
            col_count = 0
            i += 1
            continue
            
        # Default: move to next token
        i += 1

def write_docx(title, summary, dfd, schema, image_path):
    """Create a Word document with proper formatting from markdown content."""
    doc = Document()
    doc.add_heading(title, 0)
    
    # Design Summary section
    doc.add_heading('Design Summary', level=1)
    markdown_to_docx(doc, summary.content)

    # Data Flow Diagram image
    if image_path:
        doc.add_heading("Data Flow Diagram", level=1)
        doc.add_picture(image_path, width=Inches(6.0))

    # # Data Flow Diagram (Mermaid) section
    # doc.add_heading('Data Flow Diagram (Mermaid)', level=1)
    # markdown_to_docx(doc, dfd.content)

    # Database Schema section
    doc.add_heading('Database Schema (SQL)', level=1)
    markdown_to_docx(doc, schema.content)

    # Save the document
    filename = title.replace(" ", "_") + time.strftime("_%Y%m%d_%H%M%S") + ".docx"
    folder = "agent_output_design_docs"
    os.makedirs(folder, exist_ok=True)
    doc.save(os.path.join(folder, filename))
    return filename
